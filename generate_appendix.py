# -*- coding: utf-8 -*-
"""Генерация Приложений А, Б, В для ВКР Хэ Цзянь"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import json

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    return p

def set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = alignment

def shade_cells(row, color='D9E2F3'):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): color,
            qn('w:val'): 'clear'
        })
        shading.append(shd)

# ============================================================
# ПРИЛОЖЕНИЕ А
# ============================================================
doc.add_page_break()
add_para('Приложение А', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para('(справочное)', italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('Примеры ручной аннотации даньму по трём эмоциональным классам', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('')

add_para('В данном приложении представлены 18 примеров (по 6 на каждый класс) из выборки ручной '
         'аннотации (N = 430). Для каждого примера указаны: исходный китайский текст, пиньинь, '
         'русский перевод, тематический раздел видео и присвоенный класс.', size=11, italic=True)

examples = {
    'Положительные (正面)': [
        {
            'text': '那肯定',
            'pinyin': 'Nà kěndìng',
            'trans': 'Ну конечно / Естественно',
            'cat': 'Технологии',
            'note': 'Выражение согласия без эксплицитной оценочной лексики'
        },
        {
            'text': '那不得为token付费吗，还是搞个好显卡本地部署好',
            'pinyin': 'Nà bùděi wèi token fùfèi ma, háishì gǎo gè hǎo xiǎnkǎ běndì bùshǔ hǎo',
            'trans': 'Разве не придётся платить за токены? Лучше взять хорошую видеокарту и развернуть локально',
            'cat': 'Технологии',
            'note': 'Практический совет с имплицитным одобрением локального ИИ'
        },
        {
            'text': '李健这首歌比较好听',
            'pinyin': 'Lǐ Jiàn zhè shǒu gē bǐjiào hǎotīng',
            'trans': 'Эта песня Ли Цзяня довольно приятная',
            'cat': 'Развлечения',
            'note': 'Эксплицитная положительная оценка с компаративом'
        },
        {
            'text': '王一博牛逼克拉斯！！！',
            'pinyin': 'Wáng Yībó niúbī kèlāsī!!!',
            'trans': 'Ван Ибо — офигенно крут, высший класс!!!',
            'cat': 'Развлечения',
            'note': 'Фан-дискурс с молодёжным сленгом и тройным восклицанием'
        },
        {
            'text': '懂你意思',
            'pinyin': 'Dǒng nǐ yìsi',
            'trans': 'Понимаю, о чём ты / Я тебя понял',
            'cat': 'Автомобили',
            'note': 'Формула иронической солидарности (иньян-гуйци)'
        },
        {
            'text': '公路之王，只有X5，没有其他。',
            'pinyin': 'Gōnglù zhī wáng, zhǐyǒu X5, méiyǒu qítā.',
            'trans': 'Король дорог — только X5, никаких других.',
            'cat': 'Автомобили',
            'note': 'Категоричное одобрение с брендовой лояльностью'
        },
    ],
    'Нейтральные (中性)': [
        {
            'text': '不认识',
            'pinyin': 'Bù rènshi',
            'trans': 'Не знаю / Не знаком',
            'cat': 'Технологии',
            'note': 'Констатация отсутствия знания без оценки'
        },
        {
            'text': '安装在云端上',
            'pinyin': 'Ānzhuāng zài yúnduān shàng',
            'trans': 'Установлено в облаке',
            'cat': 'Технологии',
            'note': 'Описательное техническое замечание'
        },
        {
            'text': '是真的吗？',
            'pinyin': 'Shì zhēn de ma?',
            'trans': 'Это правда?',
            'cat': 'Развлечения',
            'note': 'Нейтральный запрос верификации'
        },
        {
            'text': 'Z',
            'pinyin': 'Z',
            'trans': 'Z',
            'cat': 'Развлечения',
            'note': 'Однобуквенное сообщение, смысл неопределён'
        },
        {
            'text': '冷知识：量产iPhone的水平仪都有1°左右的歪斜偏差',
            'pinyin': 'Lěng zhīshi: liàngchǎn iPhone de shuǐpíngyí dōu yǒu 1° zuǒyòu de wāixié piānchā',
            'trans': 'Холодный факт: у серийных iPhone уровень имеет отклонение около 1°',
            'cat': 'Автомобили',
            'note': 'Фактологическая реплика без эмоциональной оценки'
        },
        {
            'text': '按理来说，这个第一圈跑肯定没有第二圈跑的那么快',
            'pinyin': 'Ànlǐ láishuō, zhège dìyī quān pǎo kěndìng méiyǒu dìèr quān pǎo de nàme kuài',
            'trans': 'По логике, первый круг точно не будет таким же быстрым, как второй',
            'cat': 'Автомобили',
            'note': 'Рассудительное наблюдение без выраженной оценки'
        },
    ],
    'Отрицательные (负面)': [
        {
            'text': '坚决不玩',
            'pinyin': 'Jiānjué bù wán',
            'trans': 'Категорически не буду играть',
            'cat': 'Технологии',
            'note': 'Решительный отказ от взаимодействия с продуктом'
        },
        {
            'text': '那可太焦虑了，都给你整明白了',
            'pinyin': 'Nà kě tài jiāolǜ le, dōu gěi nǐ zhěng míngbai le',
            'trans': 'Ну это слишком тревожно, тебе всё разжевали',
            'cat': 'Технологии',
            'note': 'Саркастическое выражение тревоги по поводу ИИ'
        },
        {
            'text': '用得着她吗？博览会，科技展览会是干嘛的',
            'pinyin': 'Yòng de zháo tā ma? Bólǎnhuì, kējì zhǎnlǎnhuì shì gànmá de',
            'trans': 'Зачем она нужна? Для чего тогда выставки и технологические экспо?',
            'cat': 'Развлечения',
            'note': 'Риторический вопрос, выражающий критику'
        },
        {
            'text': '鸡皮疙瘩起来了',
            'pinyin': 'Jīpí gēda qǐlái le',
            'trans': 'Мурашки по коже пошли',
            'cat': 'Развлечения',
            'note': 'Соматическая метафора отторжения / дискомфорта'
        },
        {
            'text': '普通卡宴拉完了',
            'pinyin': 'Pǔtōng Kǎyàn lā wán le',
            'trans': 'Обычный Cayenne — полный отстой',
            'cat': 'Автомобили',
            'note': 'Жаргонное уничижение (拉 — сленг: «полный провал»)'
        },
        {
            'text': '？？？？？这么好的铺装路面那我丙察察算什么',
            'pinyin': '????? Zhème hǎo de pūzhuāng lùmiàn nà wǒ Bǐngcháchá suàn shénme',
            'trans': '????? Если это хорошая дорога, то чем тогда считать Бинчача?',
            'cat': 'Автомобили',
            'note': 'Сравнительная критика с отсылкой к легендарно плохой дороге Бинчача'
        },
    ],
}

for class_name, items in examples.items():
    add_para(class_name, bold=True, size=12)
    add_para('')

    for i, ex in enumerate(items, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f'А.{list(examples.keys()).index(class_name)+1}.{i}  ')
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(10)
        run_num.bold = True

        # Original text
        run_cn = p.add_run(f'{ex["text"]}')
        run_cn.font.name = 'Times New Roman'
        run_cn.font.size = Pt(10)

        p2 = doc.add_paragraph()
        run_py = p2.add_run(f'     Пиньинь: {ex["pinyin"]}')
        run_py.font.name = 'Times New Roman'
        run_py.font.size = Pt(10)
        run_py.italic = True

        p3 = doc.add_paragraph()
        run_tr = p3.add_run(f'     Перевод: {ex["trans"]}')
        run_tr.font.name = 'Times New Roman'
        run_tr.font.size = Pt(10)

        p4 = doc.add_paragraph()
        run_info = p4.add_run(f'     Раздел: {ex["cat"]}. Примечание: {ex["note"]}')
        run_info.font.name = 'Times New Roman'
        run_info.font.size = Pt(9)
        run_info.italic = True

        add_para('')

# Annotation note
add_para('Примечание — Ручная аннотация выполнена автором единолично. Оценка эмоциональной '
         'окраски производилась исключительно по тексту даньму, без привлечения видеоконтекста. '
         'При наличии сомнений предпочтение отдавалось нейтральному классу. Все примеры взяты '
         'из очищенной выборки даньму платформы Bilibili (апрель 2026 г.).', size=9, italic=True)

# ============================================================
# ПРИЛОЖЕНИЕ Б
# ============================================================
doc.add_page_break()
add_para('Приложение Б', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para('(справочное)', italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('Метаданные анализируемых видео', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('')

videos = [
    ['1', 'BV171Paz2ESZ',
     'Ai商业慢谈 / Деловые разговоры об AI',
     'Ai商业慢谈', 'Технологии', '2026-03-15', '8:58', '897', '681'],
    ['2', 'BV1ojfDBSEPv',
     '【闪客】一口气拆穿 Skill/MCP/RAG/Agent/OpenClaw 底层逻辑 /\nАнализ архитектуры Skill/MCP/RAG',
     '闪客', 'Технологии', '2026-02-03', '14:20', '1800', '998'],
    ['3', 'BV1SvZyBJENz',
     '山茶爱讲话 / Шанча говорит',
     '山茶爱讲话', 'Развлечения', '2026-02-17', '5:48', '1161', '1007'],
    ['4', 'BV1xi6kBFEmZ',
     '顶级公关？是把人当傻子吧！ /\nВысший PR: неужели людей считают дураками?',
     '—', 'Развлечения', '2026-01-31', '7:15', '1153', '986'],
    ['5', 'BV1Mpd5BFEk1',
     'Upspeed 盛嘉成 / Автомобильные гонки',
     'Upspeed盛嘉成', 'Автомобили', '2026-04-18', '13:29', '948', '459'],
    ['6', 'BV1cZXKBpENC',
     '公路之王挑战赛 / Шоссейный король: челлендж',
     '—', 'Автомобили', '2026-03-27', '16:42', '2145', '1401'],
]

table = doc.add_table(rows=8, cols=9)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['№', 'BV-код', 'Название видео', 'Автор канала', 'Категория',
           'Дата публ.', 'Длит.', 'Исх.\nданьму', 'После\nочистки']
for j, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[j], h, bold=True, size=9,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
shade_cells(table.rows[0])

for i, vid in enumerate(videos):
    row = table.rows[i + 1]
    for j, val in enumerate(vid):
        alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 4, 5, 6, 7, 8] else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(row.cells[j], val, size=9, alignment=alignment)

# Total row
total_row = table.rows[7]
set_cell_text(total_row.cells[0], '', size=9)
set_cell_text(total_row.cells[1], '', size=9)
set_cell_text(total_row.cells[2], '', size=9)
set_cell_text(total_row.cells[3], '', size=9)
set_cell_text(total_row.cells[4], 'Итого', bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(total_row.cells[5], '—', size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(total_row.cells[6], '—', size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(total_row.cells[7], '8104', bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(total_row.cells[8], '6314', bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_para('')
add_para('Примечание — Все видео опубликованы на платформе Bilibili в 2026 году. Даньму собраны '
         'через открытый API платформы (bilibili-api-python). Период сбора данных: 15–25 апреля 2026 г. '
         'Длительность указана в формате мин:сек. Исходное количество даньму — до очистки; '
         'после очистки — количество записей, использованных в анализе.', size=9, italic=True)

# ============================================================
# ПРИЛОЖЕНИЕ В
# ============================================================
doc.add_page_break()
add_para('Приложение В', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para('(справочное)', italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('Параметры обучения модели Chinese-BERT-wwm', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('')

# Table B.1: Model config
add_para('Таблица В.1 — Конфигурация обучения', bold=True, size=11, italic=True)
add_para('')

params = [
    ['Базовая модель', 'Chinese-BERT-wwm-ext (HFL, Cui et al., 2021)'],
    ['Архитектура классификации', 'BERT + линейный слой (3 класса)'],
    ['Размер батча (batch size)', '16'],
    ['Количество эпох', '5 (лучшая — эпоха 4)'],
    ['Максимальная длина токенов (max_len)', '64'],
    ['Скорость обучения (learning rate)', '2×10⁻⁵ (AdamW)'],
    ['Функция потерь', 'Cross-entropy loss'],
    ['Случайное зерно (random seed)', '42'],
    ['Устройство обучения', 'CPU (Intel Core), ~40 мин.'],
    ['Размер обучающей / тестовой выборки', '344 / 86 (случайное разделение 8:2)'],
    ['Общий объём размеченных данных', '430 даньму'],
    ['Фреймворк', 'HuggingFace Transformers + PyTorch'],
]

table2 = doc.add_table(rows=len(params)+1, cols=2)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

set_cell_text(table2.rows[0].cells[0], 'Параметр', bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(table2.rows[0].cells[1], 'Значение', bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
shade_cells(table2.rows[0])

for i, (param, val) in enumerate(params):
    set_cell_text(table2.rows[i+1].cells[0], param, size=10)
    set_cell_text(table2.rows[i+1].cells[1], val, size=10)

add_para('')

# Table B.2: Training curve
add_para('Таблица В.2 — Динамика обучения по эпохам', bold=True, size=11, italic=True)
add_para('')

epochs = [
    ['1', '0,9642', '58,14%', '—'],
    ['2', '0,7770', '63,95%', '—'],
    ['3', '0,6126', '65,12%', '—'],
    ['4', '0,4437', '70,93%', 'Лучшая эпоха — выбрана для прогнозирования'],
    ['5', '0,3729', '69,77%', 'Начало переобучения'],
]

table3 = doc.add_table(rows=len(epochs)+1, cols=4)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

curve_headers = ['Эпоха', 'Loss (обучение)', 'Accuracy (тест)', 'Примечание']
for j, h in enumerate(curve_headers):
    set_cell_text(table3.rows[0].cells[j], h, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
shade_cells(table3.rows[0])

for i, row_data in enumerate(epochs):
    for j, val in enumerate(row_data):
        set_cell_text(table3.rows[i+1].cells[j], val, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER if j < 3 else WD_ALIGN_PARAGRAPH.LEFT)

add_para('')
add_para('Примечание — Точность указана для тестовой выборки (n = 86). На эпохе 5 наблюдается '
         'снижение точности при продолжающемся уменьшении функции потерь на обучающей выборке, '
         'что является классическим признаком переобучения. Выбрана эпоха 4 как обеспечивающая '
         'наилучший баланс между качеством классификации и обобщающей способностью модели.', size=9, italic=True)

# Save
output_path = 'C:/Users/SOLO/WorkBuddy/20260326235845/Приложения_А_Б_В.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
